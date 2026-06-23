from __future__ import annotations

import base64
import io
import json
import os
import re
import tempfile
from dataclasses import dataclass, field
from datetime import date
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Iterable

import pdfplumber
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

APP_TITLE = "보건대행기관 업무실적 결과보고서 생성기"
MAX_UPLOAD_BYTES = 800 * 1024 * 1024
REPORT_FILE_NAME = "보건대행기관_업무실적_결과보고서.docx"
ROLE_PATTERNS = {
    "doctor": {"label": "의사", "patterns": [r"의\s*사", r"직업환경의학", r"산업보건의"]},
    "hygienist": {"label": "산업위생관리기사", "patterns": [r"산업\s*위생\s*관리\s*기사", r"산업위생", r"위생관리기사"]},
    "nurse": {"label": "간호사", "patterns": [r"간\s*호\s*사", r"보건관리\s*간호"]},
}


@dataclass
class RoleStats:
    visits: int = 0
    counseling_people: int = 0


@dataclass
class PdfResult:
    file_name: str
    institution: str
    roles: dict[str, RoleStats] = field(default_factory=lambda: {
        "doctor": RoleStats(),
        "hygienist": RoleStats(),
        "nurse": RoleStats(),
    })
    measurements: list[str] = field(default_factory=list)
    actions: list[str] = field(default_factory=list)
    materials: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    text_preview: str = ""

    def to_dict(self) -> dict:
        return {
            "fileName": self.file_name,
            "institution": self.institution,
            "doctorVisits": self.roles["doctor"].visits,
            "doctorCounseling": self.roles["doctor"].counseling_people,
            "hygienistVisits": self.roles["hygienist"].visits,
            "hygienistCounseling": self.roles["hygienist"].counseling_people,
            "nurseVisits": self.roles["nurse"].visits,
            "nurseCounseling": self.roles["nurse"].counseling_people,
            "measurements": " / ".join(self.measurements),
            "actions": " / ".join(self.actions),
            "materials": " / ".join(self.materials),
            "warnings": self.warnings,
            "textPreview": self.text_preview,
        }


def normalize_text(text: str) -> str:
    text = text.replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf_text(path: Path) -> tuple[str, list[str]]:
    warnings: list[str] = []
    chunks: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text(x_tolerance=1.5, y_tolerance=3) or ""
                if page_text.strip():
                    chunks.append(page_text)
                else:
                    warnings.append(f"{page_number}쪽 텍스트 추출 없음")
    except Exception as exc:
        warnings.append(f"PDF 읽기 실패: {exc}")
    text = normalize_text("\n".join(chunks))
    if not text:
        warnings.append("스캔 이미지 PDF일 수 있음: OCR 처리가 필요합니다.")
    return text, warnings


def collect_pdf_files(uploaded_files: list[tuple[str, bytes]], work_dir: Path) -> tuple[list[Path], list[str]]:
    pdfs: list[Path] = []
    warnings: list[str] = []
    for file_name, data in uploaded_files:
        safe_name = re.sub(r'[<>:"/\\|?*]+', "_", Path(file_name).name)
        if Path(file_name).suffix.lower() != ".pdf":
            warnings.append(f"{file_name}: PDF 파일만 처리합니다.")
            continue
        saved_path = work_dir / safe_name
        saved_path.write_bytes(data)
        pdfs.append(saved_path)
    return pdfs, warnings


def clean_field(value: str) -> str:
    value = re.split(r"\s{2,}|담당자|작성일|방문일|소재지|전화|대표자", value)[0]
    value = re.sub(r"^[\s:：\-]+|[\s:：\-]+$", "", value)
    return value[:60]


def infer_institution(file_name: str, text: str) -> str:
    patterns = [
        r"(?:사업장명|기관명|업체명|회사명)\s*[:：]?\s*([^\n]{2,40})",
        r"(?:방문\s*사업장|대상\s*사업장)\s*[:：]?\s*([^\n]{2,40})",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return clean_field(match.group(1))
    stem = Path(file_name).stem
    stem = re.sub(r"(방문보고서|보고서|의사|간호사|산업위생관리기사|산업위생|보건대행)", "", stem)
    stem = re.sub(r"[_\-\[\]\(\)\d]{2,}", " ", stem)
    return clean_field(stem) or Path(file_name).stem


def split_windows(text: str, role_regex: str, radius: int = 520) -> Iterable[str]:
    for match in re.finditer(role_regex, text, flags=re.IGNORECASE):
        yield text[max(0, match.start() - radius):min(len(text), match.end() + radius)]


def find_nearby_numbers(text: str, labels: list[str], default_one_when_role_seen: bool = False) -> int:
    values: list[int] = []
    label_expr = "|".join(labels)
    patterns = [
        rf"(?:{label_expr})\s*[:：]?\s*(\d{{1,4}})\s*(?:회|명|건|인)?",
        rf"(\d{{1,4}})\s*(?:회|명|건|인)\s*(?:{label_expr})",
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.IGNORECASE):
            values.append(int(match.group(1)))
    if values:
        return max(values)
    return 1 if default_one_when_role_seen else 0


def parse_role_stats(text: str) -> dict[str, RoleStats]:
    result = {key: RoleStats() for key in ROLE_PATTERNS}
    for key, meta in ROLE_PATTERNS.items():
        windows = list(split_windows(text, "|".join(meta["patterns"])))
        if not windows:
            continue
        role_text = "\n".join(windows)
        result[key].visits = find_nearby_numbers(
            role_text,
            [r"방문\s*횟수", r"방문\s*건수", r"방문", r"실시\s*횟수", r"횟수"],
            True,
        )
        result[key].counseling_people = find_nearby_numbers(
            role_text,
            [r"상담\s*인원", r"상담자\s*수", r"상담", r"면담\s*인원", r"교육\s*인원"],
        )
    return result


def extract_section_lines(text: str, keywords: list[str], max_items: int = 8) -> list[str]:
    lines = [line.strip(" -\t") for line in text.splitlines() if line.strip()]
    found: list[str] = []
    keyword_expr = re.compile("|".join(keywords), re.IGNORECASE)
    for index, line in enumerate(lines):
        if not keyword_expr.search(line):
            continue
        value = line
        if len(value) < 12 and index + 1 < len(lines):
            value = f"{value}: {lines[index + 1]}"
        value = re.sub(r"\s+", " ", value)
        if value not in found:
            found.append(value[:180])
        if len(found) >= max_items:
            break
    return found


def parse_pdf(path: Path) -> PdfResult:
    text, warnings = extract_pdf_text(path)
    result = PdfResult(path.name, infer_institution(path.name, text), warnings=warnings, text_preview=text[:700])
    result.roles = parse_role_stats(text)
    result.measurements = extract_section_lines(text, [r"측정", r"작업환경", r"소음", r"분진", r"유기화합물", r"유해인자"], 5)
    result.actions = extract_section_lines(text, [r"조치\s*사항", r"개선\s*사항", r"지도\s*사항", r"권고\s*사항", r"사후\s*관리"], 8)
    result.materials = extract_section_lines(text, [r"홍보물", r"자료\s*배포", r"교육\s*자료", r"안내문", r"리플릿"], 5)
    return result


def summarize(results: list[PdfResult]) -> dict:
    summary = {
        "pdfCount": len(results),
        "doctorVisits": sum(item.roles["doctor"].visits for item in results),
        "doctorCounseling": sum(item.roles["doctor"].counseling_people for item in results),
        "hygienistVisits": sum(item.roles["hygienist"].visits for item in results),
        "hygienistCounseling": sum(item.roles["hygienist"].counseling_people for item in results),
        "nurseVisits": sum(item.roles["nurse"].visits for item in results),
        "nurseCounseling": sum(item.roles["nurse"].counseling_people for item in results),
    }
    summary["totalVisits"] = summary["doctorVisits"] + summary["hygienistVisits"] + summary["nurseVisits"]
    summary["totalCounseling"] = summary["doctorCounseling"] + summary["hygienistCounseling"] + summary["nurseCounseling"]
    return summary


def unique_compact(items: Iterable[str], limit: int = 10) -> list[str]:
    result: list[str] = []
    for item in items:
        item = re.sub(r"\s+", " ", item).strip()
        if item and item not in result:
            result.append(item)
        if len(result) >= limit:
            break
    return result


def parse_multipart(content_type: str, body: bytes) -> list[tuple[str, bytes]]:
    boundary_match = re.search(r"boundary=(.+)", content_type)
    if not boundary_match:
        return []
    boundary = boundary_match.group(1).strip().strip('"').encode()
    files: list[tuple[str, bytes]] = []
    for part in body.split(b"--" + boundary):
        if b"Content-Disposition" not in part:
            continue
        header, _, data = part.partition(b"\r\n\r\n")
        if not data:
            continue
        disposition = header.decode("utf-8", errors="ignore")
        filename_match = re.search(r'filename="([^"]*)"', disposition)
        if not filename_match or not filename_match.group(1):
            continue
        files.append((Path(filename_match.group(1)).name, data.rsplit(b"\r\n", 1)[0]))
    return files


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_kv_row(table, label: str, value: str) -> None:
    cells = table.add_row().cells
    set_cell_text(cells[0], label, bold=True, size=9)
    set_cell_shading(cells[0], "F2F4F7")
    set_cell_text(cells[1], value or "-", size=9)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(9)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(2)
    for style_name, size, color in [("Heading 1", 13, "1F4D78"), ("Heading 2", 11, "1F4D78")]:
        style = doc.styles[style_name]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(4)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(18)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(subtitle)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_summary_table(doc: Document, summary: dict) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(["구분", "방문건수", "상담인원", "주요 확인사항", "비고"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9)
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
    rows = [
        ("의사", summary["doctorVisits"], summary["doctorCounseling"], "건강상담 및 의학적 관리", ""),
        ("산업위생관리기사", summary["hygienistVisits"], summary["hygienistCounseling"], "작업환경 및 유해인자 관리", ""),
        ("간호사", summary["nurseVisits"], summary["nurseCounseling"], "건강관리 상담 및 교육", ""),
        ("합계", summary["totalVisits"], summary["totalCounseling"], "", ""),
    ]
    for label, visits, counseling, check, note in rows:
        cells = table.add_row().cells
        for idx, value in enumerate([label, str(visits), str(counseling), check, note]):
            set_cell_text(cells[idx], value, bold=(label == "합계"), size=9)
        if label == "합계":
            for cell in cells:
                set_cell_shading(cell, "F2F4F7")


def add_public_report_doc(results: list[PdfResult], upload_warnings: list[str]) -> bytes:
    summary = summarize(results)
    rows = [item.to_dict() for item in results]
    actions = unique_compact((action for item in results for action in item.actions), 10)
    measurements = unique_compact((item for result in results for item in result.measurements), 8)
    materials = unique_compact((item for result in results for item in result.materials), 8)
    warnings = upload_warnings + [warning for item in results for warning in item.warnings]

    doc = Document()
    style_document(doc)
    add_title(doc, "보건대행기관 업무실적 결과보고서", f"작성일: {date.today().isoformat()} / 분석 PDF: {summary['pdfCount']}건")
    doc.add_heading("1. 보건대행기관 업무 수행 실적", level=1)
    meta = doc.add_table(rows=0, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    add_kv_row(meta, "보고기간", "업로드한 방문보고서 기준")
    add_kv_row(meta, "대상기관", f"{len(set(row['institution'] for row in rows if row['institution']))}개 기관")
    add_kv_row(meta, "근거자료", f"PDF 방문보고서 {summary['pdfCount']}건")
    doc.add_paragraph("")
    add_summary_table(doc, summary)

    doc.add_heading("2. 측정 결과 및 주요 조치사항", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(["구분", "추출 내용", "보고서 반영"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9)
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
    for label, values, note in [
        ("측정 결과", measurements, "작업환경 및 유해인자 관련 확인내용 반영"),
        ("조치사항", actions, "개선, 지도, 권고, 사후관리 사항 반영"),
        ("홍보물/교육자료", materials, "자료 배포 및 교육 실시 내용 반영"),
    ]:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, size=9)
        set_cell_text(cells[1], "\n".join(values) if values else "해당 문구 자동 추출 없음", size=8)
        set_cell_text(cells[2], note, size=8)
    note = doc.add_paragraph()
    note.add_run("※ 본 보고서는 PDF 텍스트 자동 추출 결과를 기반으로 작성되었으며, 제출 전 원문과 대조 확인이 필요합니다.").font.size = Pt(8)

    doc.add_page_break()
    doc.add_heading("3. 기관별 확인 요약", level=1)
    detail = doc.add_table(rows=1, cols=7)
    detail.alignment = WD_TABLE_ALIGNMENT.CENTER
    detail.style = "Table Grid"
    for idx, header in enumerate(["기관명", "의사", "산업위생", "간호사", "상담합계", "조치사항", "확인 필요"]):
        set_cell_text(detail.rows[0].cells[idx], header, bold=True, size=8)
        set_cell_shading(detail.rows[0].cells[idx], "E8EEF5")
    grouped: dict[str, dict] = {}
    for item in rows:
        institution = item["institution"] or "기관명 미확인"
        grouped.setdefault(institution, {"doctor": 0, "hygienist": 0, "nurse": 0, "counseling": 0, "actions": [], "warnings": []})
        group = grouped[institution]
        group["doctor"] += item["doctorVisits"]
        group["hygienist"] += item["hygienistVisits"]
        group["nurse"] += item["nurseVisits"]
        group["counseling"] += item["doctorCounseling"] + item["hygienistCounseling"] + item["nurseCounseling"]
        if item["actions"]:
            group["actions"].append(item["actions"])
        group["warnings"].extend(item["warnings"])
    for institution, group in list(grouped.items())[:18]:
        cells = detail.add_row().cells
        values = [
            institution,
            str(group["doctor"]),
            str(group["hygienist"]),
            str(group["nurse"]),
            str(group["counseling"]),
            " / ".join(unique_compact(group["actions"], 2)) or "-",
            " / ".join(unique_compact(group["warnings"], 2)) or "-",
        ]
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value[:160], size=7 if idx in {5, 6} else 8)
    if len(grouped) > 18:
        paragraph = doc.add_paragraph(f"※ 기관별 요약은 지면 관계상 18개 기관까지만 표시했습니다. 전체 분석 대상: {len(grouped)}개 기관")
        paragraph.runs[0].font.size = Pt(8)

    doc.add_heading("4. 제출 전 확인사항", level=1)
    checks = [
        "스캔 PDF 또는 텍스트 추출 실패 페이지가 있는지 확인",
        "기관별 방문건수와 상담인원이 원문 보고서와 일치하는지 확인",
        "조치사항, 개선사항, 홍보물 배포 내용의 누락 여부 확인",
        "공공기관 내부 결재 서식의 보고기간, 부서명, 담당자명 추가 기재",
    ]
    if warnings:
        checks.append(f"자동 확인 필요 항목 {len(warnings)}건 검토")
    for item in checks:
        doc.add_paragraph(item, style="List Bullet")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class ReportHandler(BaseHTTPRequestHandler):
    server_version = "HealthReportParser/2.0"

    def do_GET(self) -> None:
        if self.path in {"/", "/index.html"}:
            self.respond_html(INDEX_HTML)
            return
        self.send_error(HTTPStatus.NOT_FOUND)

    def do_POST(self) -> None:
        if self.path == "/api/report":
            self.handle_report()
            return
        self.send_error(HTTPStatus.NOT_FOUND)

    def handle_report(self) -> None:
        content_length = int(self.headers.get("Content-Length", "0"))
        if content_length > MAX_UPLOAD_BYTES:
            self.respond_json({"error": "업로드 용량이 너무 큽니다."}, HTTPStatus.REQUEST_ENTITY_TOO_LARGE)
            return
        files = parse_multipart(self.headers.get("Content-Type", ""), self.rfile.read(content_length))
        if not files:
            self.respond_json({"error": "PDF 파일을 선택해 주세요."}, HTTPStatus.BAD_REQUEST)
            return
        with tempfile.TemporaryDirectory(prefix="health_reports_") as temp_dir_name:
            pdfs, warnings = collect_pdf_files(files, Path(temp_dir_name))
            if not pdfs:
                self.respond_json({"error": "처리할 PDF 파일이 없습니다."}, HTTPStatus.BAD_REQUEST)
                return
            results = [parse_pdf(path) for path in pdfs]
            report_bytes = add_public_report_doc(results, warnings)
            self.respond_json({
                "summary": summarize(results),
                "rows": [item.to_dict() for item in results],
                "warnings": warnings,
                "reportFileName": REPORT_FILE_NAME,
                "reportBase64": base64.b64encode(report_bytes).decode("ascii"),
            })

    def respond_html(self, content: str) -> None:
        data = content.encode("utf-8")
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def respond_json(self, payload: dict, status: HTTPStatus = HTTPStatus.OK) -> None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def log_message(self, format: str, *args: object) -> None:
        print(f"{self.address_string()} - {format % args}")


INDEX_HTML = r"""<!doctype html>
<html lang="ko">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>보건대행기관 업무실적 결과보고서 생성기</title>
  <style>
    :root { --bg:#f5f7fb; --panel:#fff; --ink:#162033; --muted:#667085; --line:#d8dee9; --primary:#155eef; --primary-dark:#0f45b8; --warn:#b54708; --ok:#067647; }
    * { box-sizing:border-box; } body { margin:0; font-family:"Malgun Gothic","Apple SD Gothic Neo",system-ui,sans-serif; background:var(--bg); color:var(--ink); }
    header { background:linear-gradient(135deg,#17324d 0%,#155eef 100%); color:white; padding:34px 28px; } header .wrap, main { max-width:1280px; margin:0 auto; }
    h1 { margin:0 0 10px; font-size:30px; letter-spacing:0; } header p { max-width:900px; margin:0; line-height:1.6; color:#e5edff; } main { padding:24px; }
    .toolbar { display:grid; grid-template-columns:minmax(280px,1fr) auto; gap:16px; align-items:stretch; margin-bottom:20px; }
    .upload { border:1px dashed #98a2b3; background:var(--panel); border-radius:8px; padding:18px; min-height:118px; display:flex; flex-direction:column; justify-content:center; gap:10px; }
    .upload strong { font-size:17px; } input[type=file] { width:100%; } button { border:0; border-radius:8px; background:var(--primary); color:white; font-weight:700; padding:0 22px; min-height:52px; cursor:pointer; font-size:15px; }
    button:hover { background:var(--primary-dark); } button:disabled { opacity:.55; cursor:wait; } .summary { display:grid; grid-template-columns:repeat(4,minmax(150px,1fr)); gap:12px; margin-bottom:20px; }
    .metric { background:var(--panel); border:1px solid var(--line); border-radius:8px; padding:15px; } .metric span { color:var(--muted); font-size:13px; } .metric b { display:block; font-size:26px; margin-top:5px; }
    .panel { background:var(--panel); border:1px solid var(--line); border-radius:8px; overflow:hidden; } .panel-head { padding:14px 16px; border-bottom:1px solid var(--line); display:flex; align-items:center; justify-content:space-between; gap:12px; }
    .panel-head h2 { font-size:18px; margin:0; } .secondary { min-height:38px; padding:0 14px; background:#344054; } .table-wrap { overflow:auto; max-height:58vh; }
    table { width:100%; border-collapse:collapse; min-width:1120px; } th, td { border-bottom:1px solid var(--line); padding:10px 12px; text-align:left; vertical-align:top; font-size:13px; }
    th { position:sticky; top:0; background:#f8fafc; z-index:1; white-space:nowrap; } td.num { text-align:right; font-variant-numeric:tabular-nums; } .empty { padding:46px 20px; text-align:center; color:var(--muted); }
    .notice { margin:14px 0; padding:12px 14px; border-radius:8px; border:1px solid #fedf89; background:#fffaeb; color:var(--warn); display:none; white-space:pre-wrap; } .ok { color:var(--ok); font-weight:700; }
    @media (max-width:800px) { header{padding:26px 18px;} main{padding:16px;} .toolbar{grid-template-columns:1fr;} .summary{grid-template-columns:repeat(2,minmax(0,1fr));} button{width:100%;} }
  </style>
</head>
<body>
  <header><div class="wrap"><h1>보건대행기관 업무실적 결과보고서 생성기</h1><p>의사, 산업위생관리기사, 간호사가 제출한 PDF 방문보고서를 여러 개 올리면 방문건수, 상담인원, 측정 결과, 조치사항, 홍보물 배포 내용을 추출해 2쪽 구성의 결과보고서 문서로 만듭니다.</p></div></header>
  <main>
    <section class="toolbar"><label class="upload"><strong>PDF 파일만 선택</strong><span>여러 PDF를 한 번에 선택할 수 있습니다. 스캔본 PDF는 텍스트 추출이 안 될 수 있습니다.</span><input id="files" type="file" multiple accept=".pdf,application/pdf"></label><button id="analyze">결과보고서 만들기</button></section>
    <div id="notice" class="notice"></div>
    <section class="summary" aria-label="집계 요약"><div class="metric"><span>분석 PDF</span><b id="pdfCount">0</b></div><div class="metric"><span>총 방문건수</span><b id="totalVisits">0</b></div><div class="metric"><span>총 상담인원</span><b id="totalCounseling">0</b></div><div class="metric"><span>처리 상태</span><b id="statusText">대기</b></div></section>
    <section class="panel"><div class="panel-head"><h2>문서 생성 전 미리보기</h2><button id="download" class="secondary" disabled>보고서 다운로드</button></div><div class="table-wrap"><table><thead><tr><th>기관명</th><th>파일명</th><th>의사 방문</th><th>의사 상담</th><th>산업위생 방문</th><th>산업위생 상담</th><th>간호사 방문</th><th>간호사 상담</th><th>측정/조치사항</th><th>확인 필요</th></tr></thead><tbody id="rows"><tr><td colspan="10" class="empty">아직 분석한 PDF가 없습니다.</td></tr></tbody></table></div></section>
  </main>
  <script>
    const files = document.querySelector("#files"), analyze = document.querySelector("#analyze"), download = document.querySelector("#download"), notice = document.querySelector("#notice"), rows = document.querySelector("#rows");
    let reportBlob = null, reportFileName = "보건대행기관_업무실적_결과보고서.docx";
    const esc = (value) => String(value ?? "").replace(/[&<>"']/g, (char) => ({"&":"&amp;","<":"&lt;",">":"&gt;",'"':"&quot;","'":"&#39;"}[char]));
    function setMetric(id, value) { document.querySelector("#" + id).textContent = value ?? 0; }
    function showNotice(messages) { const text = Array.isArray(messages) ? messages.join("\n") : messages; notice.textContent = text || ""; notice.style.display = text ? "block" : "none"; }
    function base64ToBlob(base64) { const bytes = atob(base64); const chunks = []; for (let i = 0; i < bytes.length; i += 8192) { const slice = bytes.slice(i, i + 8192); const array = new Uint8Array(slice.length); for (let j = 0; j < slice.length; j++) array[j] = slice.charCodeAt(j); chunks.push(array); } return new Blob(chunks, { type:"application/vnd.openxmlformats-officedocument.wordprocessingml.document" }); }
    function render(payload) { const summary = payload.summary || {}; setMetric("pdfCount", summary.pdfCount); setMetric("totalVisits", summary.totalVisits); setMetric("totalCounseling", summary.totalCounseling); document.querySelector("#statusText").innerHTML = '<span class="ok">완료</span>'; const data = payload.rows || []; rows.innerHTML = data.length ? data.map((row) => `<tr><td>${esc(row.institution)}</td><td>${esc(row.fileName)}</td><td class="num">${esc(row.doctorVisits)}</td><td class="num">${esc(row.doctorCounseling)}</td><td class="num">${esc(row.hygienistVisits)}</td><td class="num">${esc(row.hygienistCounseling)}</td><td class="num">${esc(row.nurseVisits)}</td><td class="num">${esc(row.nurseCounseling)}</td><td>${esc([row.measurements,row.actions,row.materials].filter(Boolean).join(" / "))}</td><td>${esc((row.warnings || []).join(" / "))}</td></tr>`).join("") : '<tr><td colspan="10" class="empty">PDF를 찾지 못했습니다.</td></tr>'; reportBlob = base64ToBlob(payload.reportBase64 || ""); reportFileName = payload.reportFileName || reportFileName; download.disabled = !payload.reportBase64; showNotice(payload.warnings || ""); }
    analyze.addEventListener("click", async () => { if (!files.files.length) { showNotice("분석할 PDF 파일을 선택해 주세요."); return; } const form = new FormData(); for (const file of files.files) form.append("files", file); analyze.disabled = true; download.disabled = true; reportBlob = null; document.querySelector("#statusText").textContent = "분석 중"; showNotice(""); try { const response = await fetch("/api/report", { method:"POST", body:form }); const payload = await response.json(); if (!response.ok) throw new Error(payload.error || "보고서 생성 중 오류가 발생했습니다."); render(payload); } catch (error) { document.querySelector("#statusText").textContent = "오류"; showNotice(error.message); } finally { analyze.disabled = false; } });
    download.addEventListener("click", () => { if (!reportBlob) return; const url = URL.createObjectURL(reportBlob); const a = document.createElement("a"); a.href = url; a.download = reportFileName; document.body.appendChild(a); a.click(); a.remove(); URL.revokeObjectURL(url); });
  </script>
</body>
</html>
"""


def main() -> None:
    port = int(os.environ.get("PORT", "8000"))
    server = ThreadingHTTPServer(("127.0.0.1", port), ReportHandler)
    print(f"{APP_TITLE} 실행 중: http://127.0.0.1:{port}")
    print("종료하려면 Ctrl+C를 누르세요.")
    server.serve_forever()


if __name__ == "__main__":
    main()
